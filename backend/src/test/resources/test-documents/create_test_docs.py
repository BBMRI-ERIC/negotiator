#!/usr/bin/env python3
"""
Script to create minimal valid DOC and DOCX files for testing.
This ensures we have proper test documents that can be converted to PDF.
"""

import os
import zipfile
from docx import Document
from docx.shared import Inches

def create_test_docx():
    """Create a minimal valid DOCX file."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test DOCX document for PDF conversion testing.')
    doc.add_paragraph('It contains minimal content to ensure proper conversion.')
    
    # Add a simple table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Test'
    table.cell(1, 1).text = 'Document'
    
    doc.save('test.docx')
    print("Created test.docx")

def create_simple_docx():
    """Create a very simple DOCX file manually."""
    # Create a minimal DOCX structure
    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''

    app_props = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
<Application>Test</Application>
<DocSecurity>0</DocSecurity>
<ScaleCrop>false</ScaleCrop>
<SharedDoc>false</SharedDoc>
<HyperlinksChanged>false</HyperlinksChanged>
<AppVersion>1.0</AppVersion>
</Properties>'''

    main_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''

    word_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''

    document_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p>
<w:r>
<w:t>Test Document Content</w:t>
</w:r>
</w:p>
<w:p>
<w:r>
<w:t>This is a test document for PDF conversion.</w:t>
</w:r>
</w:p>
</w:body>
</w:document>'''

    core_props = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<dc:title>Test Document</dc:title>
<dc:creator>Test</dc:creator>
<dcterms:created xsi:type="dcterms:W3CDTF">2023-01-01T12:00:00Z</dcterms:created>
<dcterms:modified xsi:type="dcterms:W3CDTF">2023-01-01T12:00:00Z</dcterms:modified>
</cp:coreProperties>'''

    with zipfile.ZipFile('simple_test.docx', 'w', zipfile.ZIP_DEFLATED) as docx:
        docx.writestr('[Content_Types].xml', content_types)
        docx.writestr('_rels/.rels', main_rels)
        docx.writestr('word/_rels/document.xml.rels', word_rels)
        docx.writestr('word/document.xml', document_xml)
        docx.writestr('docProps/app.xml', app_props)
        docx.writestr('docProps/core.xml', core_props)
    
    print("Created simple_test.docx")

if __name__ == "__main__":
    try:
        create_test_docx()
    except ImportError:
        print("python-docx not available, creating simple DOCX manually")
        create_simple_docx()